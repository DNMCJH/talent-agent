"""Resume export: convert generated resume text to downloadable DOCX."""

from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.core.deps import get_current_user
from app.models.user import User

router = APIRouter()


class ExportResumeIn(BaseModel):
    name: str = "Resume"
    sections: list[dict]  # [{title: str, content: str}]


@router.post("/export/docx")
async def export_resume_docx(
    body: ExportResumeIn,
    user: User = Depends(get_current_user),
) -> StreamingResponse:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading(body.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for sec in body.sections:
        doc.add_heading(sec.get("title", ""), level=1)
        content = sec.get("content", "")
        for para_text in content.split("\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                if para_text.strip().startswith("- ") or para_text.strip().startswith("• "):
                    p.style = "List Bullet"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{body.name.replace(' ', '_')}_resume.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
