"""Parse resume files (PDF/DOCX) into structured data via text extraction + LLM."""

from __future__ import annotations

import asyncio
import io
import json

from pydantic import BaseModel

from app.core.llm import call_llm


class ParsedResume(BaseModel):
    name: str = ""
    email: str = ""
    phone: str = ""
    education: list[dict[str, str]] = []
    experience: list[dict[str, str]] = []
    skills: list[str] = []
    projects: list[dict[str, str]] = []
    raw_text: str = ""


def extract_text_from_pdf(content: bytes) -> str:
    import fitz  # pymupdf

    doc = fitz.open(stream=content, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n".join(pages)


def extract_text_from_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]

    # Table-layout resumes — common in Chinese templates — put nearly all
    # content inside tables, which doc.paragraphs does not cover. Without this
    # such resumes parse out near-empty.
    for table in doc.tables:
        for row in table.rows:
            # row.cells repeats merged cells; dedupe while keeping order.
            seen: list[str] = []
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in seen:
                    seen.append(text)
            if seen:
                parts.append(" | ".join(seen))

    return "\n".join(parts)


_PARSE_SYSTEM = """You are a resume parser. Extract structured information from the resume text the user provides.

Return a JSON object with these fields:
- name: candidate's full name
- email: email address (empty string if not found)
- phone: phone number (empty string if not found)
- education: array of objects with {school, degree, major, period}
- experience: array of objects with {company, role, period, description}
- skills: array of skill strings
- projects: array of objects with {name, description, tech_stack, highlights}

Return ONLY valid JSON, no markdown fences."""

_RETRY_HINT = "\n\nIMPORTANT: your previous reply was not valid JSON. Reply with ONLY a single JSON object, starting with { and ending with }."


def _extract_json(response: str) -> dict:
    """Parse an LLM reply into a dict, tolerating markdown fences. {} on failure."""
    try:
        return json.loads(response)
    except json.JSONDecodeError:
        pass
    cleaned = response.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("\n", 1)[-1].rsplit("```", 1)[0]
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return {}


async def parse_resume(content: bytes, filename: str) -> ParsedResume:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # PDF/DOCX extraction is synchronous CPU/IO — offload so a large file does
    # not block the event loop (production runs a single uvicorn worker).
    if ext == "pdf":
        raw_text = await asyncio.to_thread(extract_text_from_pdf, content)
    elif ext in ("docx", "doc"):
        raw_text = await asyncio.to_thread(extract_text_from_docx, content)
    else:
        raw_text = content.decode("utf-8", errors="replace")

    if not raw_text.strip():
        return ParsedResume(raw_text="(empty document)")

    # 12k chars comfortably covers a 2-3 page resume — including table-layout
    # ones whose cell joins inflate the character count.
    truncated = raw_text[:12000]
    response = await call_llm(_PARSE_SYSTEM, truncated)
    data = _extract_json(response)

    # One retry: a malformed reply (empty dict) is usually a stray prose
    # preamble, not a content problem — re-ask with a stricter instruction.
    if not data:
        response = await call_llm(_PARSE_SYSTEM + _RETRY_HINT, truncated)
        data = _extract_json(response)

    return ParsedResume(
        name=data.get("name", ""),
        email=data.get("email", ""),
        phone=data.get("phone", ""),
        education=data.get("education", []),
        experience=data.get("experience", []),
        skills=data.get("skills", []),
        projects=data.get("projects", []),
        raw_text=raw_text[:6000],
    )
